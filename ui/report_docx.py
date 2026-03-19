"""
DOCX Report Generator for Jyotisha App
Uses python-docx to produce a detailed, narrative-style Word document.
"""

import io
import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GOLD    = RGBColor(0xC8, 0x90, 0x10)
PURPLE  = RGBColor(0x70, 0x40, 0xB0)
LIGHT   = RGBColor(0xE8, 0xD8, 0xA8)
MUTED   = RGBColor(0x80, 0x80, 0xA0)

SIGN_GLYPHS = ["♈","♉","♊","♋","♌","♍","♎","♏","♐","♑","♒","♓"]
PLANET_GLYPH = {
    "Surya":"☀","Chandra":"🌙","Budha":"☿","Shukra":"♀","Mangal":"♂",
    "Guru":"♃","Shani":"♄","Rahu":"☊","Ketu":"☋","Lagna":"↑"
}
AREA_LABEL = {
    "career":"Career & Status","health":"Health & Vitality",
    "relationships":"Relationships & Family","finances":"Finances & Wealth",
    "spirituality":"Spirituality & Upaya",
}
CAT_COLOR = {
    "Outstanding": RGBColor(0xC8, 0x90, 0x10),
    "Excellent":   RGBColor(0x20, 0xA0, 0x40),
    "Very Good":   RGBColor(0x20, 0xA0, 0xA0),
    "Good":        RGBColor(0x20, 0x60, 0xC8),
    "Moderate":    RGBColor(0xD0, 0x78, 0x00),
    "Challenging": RGBColor(0xC0, 0x20, 0x20),
    "Difficult":   RGBColor(0x90, 0x10, 0x10),
}

# ── Low-level helpers ─────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def _cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top",top),("bottom",bottom),("left",left),("right",right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def _heading(doc, text, level=1):
    p  = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = GOLD
    if level == 1:
        run.font.size = Pt(16)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(8)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "C89010")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0xA0, 0x78, 0x20)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xC0, 0x90, 0xFF)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p

def _para(doc, text, bold=False, color=None, size=11, italic=False, space_after=4):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold        = bold
    run.italic      = italic
    run.font.size   = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def _kv(doc, key, value):
    p   = doc.add_paragraph()
    k_r = p.add_run(f"{key}: ")
    k_r.bold = True
    k_r.font.color.rgb = GOLD
    k_r.font.size = Pt(11)
    v_r = p.add_run(value)
    v_r.font.size = Pt(11)
    v_r.font.color.rgb = RGBColor(0xD0, 0xC8, 0xE0)
    p.paragraph_format.space_after = Pt(2)

def _divider(doc):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2A2040")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

# ── Section builders ──────────────────────────────────────────────────────────

def _add_title_page(doc, title, subtitle, names, meta_lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ॐ")
    r.font.size = Pt(28)
    r.font.color.rgb = GOLD
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after  = Pt(12)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(title)
    r2.font.size = Pt(22)
    r2.font.bold = True
    r2.font.color.rgb = GOLD
    p2.paragraph_format.space_after = Pt(4)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    r3.font.size = Pt(10)
    r3.font.color.rgb = MUTED
    p3.paragraph_format.space_after = Pt(24)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(names)
    r4.font.size = Pt(18)
    r4.font.bold = True
    r4.font.color.rgb = LIGHT
    p4.paragraph_format.space_after = Pt(12)

    for line in meta_lines:
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rm = pm.add_run(line)
        rm.font.size = Pt(10)
        rm.font.color.rgb = MUTED
        pm.paragraph_format.space_after = Pt(2)

    today_str = datetime.date.today().strftime("%d %B %Y")
    pd = doc.add_paragraph()
    pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = pd.add_run(f"Report generated: {today_str}")
    rd.font.size = Pt(9)
    rd.font.color.rgb = RGBColor(0x50, 0x50, 0x70)
    pd.paragraph_format.space_before = Pt(20)
    _divider(doc)

def _add_overview(doc, chart, place):
    _heading(doc, "🗺️ Chart Overview", 1)
    _kv(doc, "Lagna (Ascendant)", f"{chart['lagna_san']} ({chart['lagna_en']}) — {chart['lagna_deg']:.1f}° | Lord: {chart['lagna_lord']}")
    _kv(doc, "Janma Rashi",       f"{chart['moon_san']} ({chart['moon_en']}) — {chart['moon_deg']:.1f}°")
    _kv(doc, "Janma Nakshatra",   f"{chart['nakshatra']}, Pada {chart['nak_pada']} | Lord: {chart['nak_lord']} | Deity: {chart['nak_deity']}")
    _kv(doc, "Active Maha Dasha", f"{chart['active_md']['lord']} MD — ends {chart['active_md']['end'].strftime('%b %Y')}")
    _kv(doc, "Active Antardasha", f"{chart['active_ad']['lord']} AD — ends {chart['active_ad']['end'].strftime('%b %Y')}")
    _kv(doc, "Place of Birth",    place)

def _add_planets(doc, chart):
    _heading(doc, "🪐 Planetary Positions (Nirayana Sidereal)", 1)
    _para(doc, "All positions calculated using Lahiri Ayanamsha (Nirayana / Sidereal system).", color=MUTED, size=10)

    ORDER = ["Lagna","Surya","Chandra","Budha","Shukra","Mangal","Guru","Shani","Rahu","Ketu"]
    planets = [(nm, chart["planets"].get(nm)) for nm in ORDER if chart["planets"].get(nm)]

    tbl = doc.add_table(rows=1+len(planets), cols=5)
    tbl.style = "Table Grid"
    for i, h in enumerate(["Graha","Rashi","Degree","House","Dignity"]):
        cell = tbl.rows[0].cells[i]
        _set_cell_bg(cell, "160E28")
        _cell_margins(cell)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = GOLD; r.font.size = Pt(10)

    for row_i, (nm, pd) in enumerate(planets):
        row  = tbl.rows[row_i+1]
        vals = [f"{PLANET_GLYPH.get(nm,'')} {nm}",
                f"{SIGN_GLYPHS[pd['sign']]} {pd['sign_san']}",
                f"{pd['deg']:.1f}°", str(pd['house']), pd['dignity']]
        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            _set_cell_bg(cell, "0A0C18" if row_i%2==0 else "0C0E1A")
            _cell_margins(cell)
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xC0, 0xB8, 0xD8)
            if ci == 4:
                if "Exalted" in val: r.font.color.rgb = RGBColor(0xF0,0xC0,0x40); r.bold = True
                elif "Own"   in val: r.font.color.rgb = RGBColor(0x40,0xD0,0x60); r.bold = True
                elif "Debil" in val: r.font.color.rgb = RGBColor(0xF0,0x60,0x60)
    doc.add_paragraph()

def _add_yogas(doc, chart):
    _heading(doc, "✨ Yoga Vivechana (Planetary Combinations)", 1)
    if not chart["yogas"]:
        _para(doc, "Standard chart — no classical Yogas detected. Individual planetary strengths govern the native's destiny.", color=MUTED)
        return
    for yoga in chart["yogas"]:
        _heading(doc, f"{yoga['name']}  —  ⭐ {yoga['strength']}", 2)
        _para(doc, f"{yoga['planets']} | House {yoga['house']}", color=MUTED, size=10)
        _para(doc, yoga['description'], size=11)
        doc.add_paragraph()

def _add_dasha(doc, chart):
    _heading(doc, "🗓️ Vimshottari Dasha Sequence", 1)
    today = datetime.date.today()
    md = chart["active_md"]; ad = chart["active_ad"]
    _para(doc, f"Current Maha Dasha: {md['lord']} MD  ({md['start'].strftime('%b %Y')} — {md['end'].strftime('%b %Y')})",
          bold=True, color=RGBColor(0xC0,0x90,0xFF))
    _para(doc, f"Current Antardasha: {ad['lord']} AD  ({ad['start'].strftime('%b %Y')} — {ad['end'].strftime('%b %Y')})",
          color=RGBColor(0xA0,0x70,0xE0))
    doc.add_paragraph()

    tbl = doc.add_table(rows=1+len(chart["dasha_seq"]), cols=4)
    tbl.style = "Table Grid"
    for i, h in enumerate(["Maha Dasha","Start","End","Duration"]):
        cell = tbl.rows[0].cells[i]
        _set_cell_bg(cell, "160E28"); _cell_margins(cell)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = GOLD; r.font.size = Pt(10)

    for row_i, d in enumerate(chart["dasha_seq"]):
        active = d["start"] <= today <= d["end"]
        bg     = "1A0828" if active else ("0A0C18" if row_i%2==0 else "0C0E1A")
        vals   = [f"{'▶ ' if active else ''}{d['lord']} MD",
                  d["start"].strftime("%b %Y"), d["end"].strftime("%b %Y"), f"{d['years']:.0f} yrs"]
        for ci, val in enumerate(vals):
            cell = tbl.rows[row_i+1].cells[ci]
            _set_cell_bg(cell, bg); _cell_margins(cell)
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xC0,0x90,0xFF) if active else RGBColor(0xC0,0xB8,0xD8)
            if active: r.bold = True
    doc.add_paragraph()

def _add_year_scores(doc, year_scores):
    _heading(doc, "📈 Year-Wise Forecast", 1)
    _para(doc, "Scores derived from Vimshottari Dasha quality, Jupiter and Saturn transit modifiers "
               "per Brihat Parashara Hora Shastra.", color=MUTED, size=10)
    doc.add_paragraph()

    tbl = doc.add_table(rows=1+len(year_scores), cols=5)
    tbl.style = "Table Grid"
    for i, h in enumerate(["Year","Maha Dasha","Antardasha","Score","Category"]):
        cell = tbl.rows[0].cells[i]
        _set_cell_bg(cell, "160E28"); _cell_margins(cell)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = GOLD; r.font.size = Pt(10)

    for row_i, r in enumerate(year_scores):
        row    = tbl.rows[row_i+1]
        md_str = r.get("md_lord", r.get("md1", "?"))
        ad_str = r.get("ad_lord", r.get("ad1", "?"))
        sc_val = r.get("score",   r.get("combined", 0))
        vals   = [str(r["year"]), md_str, ad_str, f"{sc_val:.1f}/10", r.get("category","")]
        cat_col = CAT_COLOR.get(r.get("category",""), RGBColor(0xC0,0xB8,0xD8))
        bg = "0A0C18" if row_i%2==0 else "0C0E1A"
        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            _set_cell_bg(cell, bg); _cell_margins(cell)
            rr = cell.paragraphs[0].add_run(val)
            rr.font.size = Pt(10)
            rr.font.color.rgb = cat_col if ci >= 3 else RGBColor(0xC0,0xB8,0xD8)
            if ci == 3: rr.bold = True
    doc.add_paragraph()

    outstanding = [r for r in year_scores if r.get("category") == "Outstanding"]
    excellent   = [r for r in year_scores if r.get("category") == "Excellent"]
    challenging = [r for r in year_scores if r.get("category") in ("Challenging","Difficult")]
    if outstanding:
        _para(doc, f"⭐ Outstanding Years: {', '.join(str(r['year']) for r in outstanding)}", bold=True, color=GOLD)
    if excellent:
        _para(doc, f"✨ Excellent Years: {', '.join(str(r['year']) for r in excellent)}", bold=True, color=RGBColor(0x40,0xD0,0x60))
    if challenging:
        _para(doc, f"⚠️ Challenging Years: {', '.join(str(r['year']) for r in challenging)}", bold=True, color=RGBColor(0xF0,0x60,0x60))

def _add_predictions(doc, predictions, name):
    _heading(doc, f"🔮 Dasha-Wise Predictions — {name}", 1)
    _para(doc, "Grounded in the Vimshottari Dasha system (BPHS Ch.46). Each Maha Dasha represents "
               "a karmic chapter. Antardasha sub-periods modulate its energy within.", color=MUTED, size=10)
    doc.add_paragraph()

    for pred in predictions:
        cat_col = CAT_COLOR.get(pred["category"], RGBColor(0x80,0x80,0xA0))
        tag     = "  ▶ ACTIVE NOW" if pred["is_current"] else ("  [Past Period]" if pred["is_past"] else "")
        glyph   = PLANET_GLYPH.get(pred["lord"], "")

        _heading(doc, f"{glyph} {pred['lord']} Maha Dasha{tag}", 2)
        p = doc.add_paragraph()
        r1 = p.add_run(f"{pred['start'].strftime('%b %Y')} — {pred['end'].strftime('%b %Y')}  ·  {pred['years']:.1f} years  ·  Score: {pred['avg_score']}/10  ·  ")
        r1.font.size = Pt(10); r1.font.color.rgb = MUTED
        r2 = p.add_run(pred["category"])
        r2.font.size = Pt(10); r2.font.bold = True; r2.font.color.rgb = cat_col
        p.paragraph_format.space_after = Pt(4)

        # Shastra ref with left border
        p2 = doc.add_paragraph()
        pPr = p2._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"),"single"); left.set(qn("w:sz"),"12")
        left.set(qn("w:space"),"8"); left.set(qn("w:color"),"7040C0")
        pBdr.append(left); pPr.append(pBdr)
        r_shas = p2.add_run(pred.get("shastra_ref",""))
        r_shas.font.italic = True; r_shas.font.size = Pt(9)
        r_shas.font.color.rgb = RGBColor(0xA0,0x80,0xC0)
        p2.paragraph_format.space_after = Pt(6)

        _para(doc, f"House {pred['house']} in {pred['sign']} — {pred['dignity']}", bold=True, size=11)
        _para(doc, pred.get("quality_note",""), color=MUTED, size=10)

        _heading(doc, "Domain-Wise Predictions", 3)
        for area, text in pred.get("life_areas", {}).items():
            label = AREA_LABEL.get(area, area.title())
            p = doc.add_paragraph()
            r_lbl = p.add_run(f"{label}: ")
            r_lbl.bold = True; r_lbl.font.size = Pt(11); r_lbl.font.color.rgb = LIGHT
            r_txt = p.add_run(text)
            r_txt.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.25)

        if pred.get("antardashas"):
            _heading(doc, "Antardasha Sub-Periods", 3)
            for ad in pred["antardashas"]:
                ac  = CAT_COLOR.get(ad["category"], RGBColor(0x80,0x80,0xA0))
                now = " ◄ ACTIVE" if ad["is_current"] else ""
                p   = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                r1 = p.add_run(f"{pred['lord']}/{ad['lord']} AD{now}  ")
                r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = ac
                r2 = p.add_run(f"({ad['start'].strftime('%b %Y')} – {ad['end'].strftime('%b %Y')})  Score: {ad['score']}/10 — {ad['category']}")
                r2.font.size = Pt(9); r2.font.color.rgb = MUTED
                p2 = doc.add_paragraph(ad.get("prediction",""))
                p2.paragraph_format.left_indent = Inches(0.4)
                p2.paragraph_format.space_after = Pt(3)
                if p2.runs: p2.runs[0].font.size = Pt(10)

        g = pred.get("guidance", {})
        if g:
            _heading(doc, "Upaya (Vedic Remedies) & Jyotishi's Counsel", 3)
            _para(doc, f"🕉️ Mantra: {g.get('mantra','')}", size=10)
            _para(doc, f"🌸 Dana (Charity): {g.get('charity','')}", size=10)
            _para(doc, f"🛕 Pilgrimage: {g.get('pilgrimage','')}", size=10)
            p = doc.add_paragraph()
            r1 = p.add_run("💫 Jyotishi's Counsel: ")
            r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = RGBColor(0xA0,0xC8,0xE0)
            r2 = p.add_run(g.get("focus",""))
            r2.font.size = Pt(11)

        _divider(doc)
        doc.add_paragraph()

def _add_combined_protection(doc, cs_all, n1, n2):
    _heading(doc, "🛡️ Dampati Protection Analysis — Year by Year", 1)
    _para(doc, "Per BPHS Ch.18: 'In a dharmic marriage, when one spouse is in karmic difficulty, "
               "the other's Dasha energy provides shelter and strength — this is the sacred geometry "
               "of Grihastha Ashrama.'", italic=True, color=MUTED, size=10)
    doc.add_paragraph()

    tbl = doc.add_table(rows=1+len(cs_all), cols=7)
    tbl.style = "Table Grid"
    for i, h in enumerate(["Year", f"{n1[:8]} Dasha", f"{n1[:6]}⭐",
                            f"{n2[:8]} Dasha", f"{n2[:6]}⭐", "Combined", "Dynamic"]):
        cell = tbl.rows[0].cells[i]
        _set_cell_bg(cell, "160E28"); _cell_margins(cell)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = GOLD; r.font.size = Pt(9)

    for row_i, r in enumerate(cs_all):
        s1, s2, comb = r["score1"], r["score2"], r["combined"]
        cat     = r["category"]
        cat_col = CAT_COLOR.get(cat, RGBColor(0xC0,0xB8,0xD8))
        if s1 >= 8.5 and s2 >= 8.5:   dynamic = "🌟 Both flourish"
        elif s1 >= 8.0 and s2 < 6.5:  dynamic = f"🛡️ {n1[:8]} protects"
        elif s2 >= 8.0 and s1 < 6.5:  dynamic = f"🛡️ {n2[:8]} protects"
        elif s1 < 5.5 and s2 < 5.5:   dynamic = "🤝 Mutual support"
        else:                           dynamic = "⚖️ Steady"
        bg   = "0A0C18" if row_i%2==0 else "0C0E1A"
        vals = [str(r["year"]), f"{r['md1']}-{r['ad1']}", f"{s1:.1f}",
                f"{r['md2']}-{r['ad2']}", f"{s2:.1f}", f"{comb:.1f}", dynamic]
        for ci, val in enumerate(vals):
            cell = tbl.rows[row_i+1].cells[ci]
            _set_cell_bg(cell, bg)
            _cell_margins(cell, top=50, bottom=50, left=80, right=80)
            rr = cell.paragraphs[0].add_run(val)
            rr.font.size = Pt(9)
            if ci in (2, 4, 5): rr.font.color.rgb = cat_col; rr.bold = True
            else: rr.font.color.rgb = RGBColor(0xC0,0xB8,0xD8)
    doc.add_paragraph()

def _make_doc():
    doc = Document()
    for section in doc.sections:
        section.page_width  = Cm(21.59)
        section.page_height = Cm(27.94)
        section.top_margin = section.bottom_margin = Cm(2.54)
        section.left_margin = section.right_margin = Cm(2.54)
    style = doc.styles['Normal']
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0xD0, 0xC8, 0xE0)
    return doc

def _to_bytes(doc):
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ── Public API ────────────────────────────────────────────────────────────────

def generate_individual_docx(chart, name, place, dob, tob_h, year_scores, predictions):
    doc = _make_doc()
    dob_str = dob.strftime("%d %B %Y") if hasattr(dob,"strftime") else str(dob)
    tob_str = f"{int(tob_h):02d}:{int((tob_h%1)*60):02d} IST"

    _add_title_page(doc,
        title="Jyotisha Vedic Astrology Report",
        subtitle="Nirayana Sidereal · Lahiri Ayanamsha · Vimshottari Dasha\nBPHS · Brihat Jataka · Phaladeepika · Saravali · Uttara Kalamrita",
        names=name,
        meta_lines=[f"Date of Birth: {dob_str}   |   Time: {tob_str}", f"Place of Birth: {place}"])
    doc.add_page_break()
    _add_overview(doc, chart, place)
    _divider(doc)
    _add_planets(doc, chart)
    doc.add_page_break()
    _add_yogas(doc, chart)
    _divider(doc)
    _add_dasha(doc, chart)
    doc.add_page_break()
    _add_year_scores(doc, year_scores)
    doc.add_page_break()
    _add_predictions(doc, predictions, name)
    _divider(doc)
    _para(doc, "ॐ तत्सत् · This report is for spiritual guidance and self-reflection. "
               "Please consult a qualified Jyotishi for major life decisions.",
          italic=True, color=MUTED, size=9)
    return _to_bytes(doc)


def generate_combined_docx(chart1, chart2, n1, n2, place1, place2,
                            dob1, dob2, tob1, tob2,
                            combined_scores, s1_data, s2_data, preds1, preds2):
    doc = _make_doc()
    dob1_str = dob1.strftime("%d %B %Y") if hasattr(dob1,"strftime") else str(dob1)
    dob2_str = dob2.strftime("%d %B %Y") if hasattr(dob2,"strftime") else str(dob2)
    t1_str   = f"{int(tob1):02d}:{int((tob1%1)*60):02d} IST"
    t2_str   = f"{int(tob2):02d}:{int((tob2%1)*60):02d} IST"

    outstanding = [r for r in combined_scores if r["combined"] >= 9.0]
    excellent   = [r for r in combined_scores if 8.0 <= r["combined"] < 9.0]
    challenging = [r for r in combined_scores if r["combined"] < 6.5]
    avg_score   = sum(r["combined"] for r in combined_scores) / max(len(combined_scores), 1)

    _add_title_page(doc,
        title="Dampati Jyotisha — Combined Report",
        subtitle="Ardhanarishvara Kundali · Vimshottari Dasha System\nBPHS Ch.18 · Jataka Parijata · Phaladeepika · Saravali",
        names=f"{n1}  &  {n2}",
        meta_lines=[f"{n1}: {dob1_str} · {t1_str} · {place1}",
                    f"{n2}: {dob2_str} · {t2_str} · {place2}"])

    doc.add_page_break()
    _heading(doc, "📊 Combined Summary", 1)
    _kv(doc, "Years Analysed",    str(len(combined_scores)))
    _kv(doc, "Average Score",     f"{avg_score:.1f}/10")
    _kv(doc, "Outstanding Years", ", ".join(str(r["year"]) for r in outstanding) or "—")
    _kv(doc, "Excellent Years",   ", ".join(str(r["year"]) for r in excellent)   or "—")
    _kv(doc, "Challenging Years", ", ".join(str(r["year"]) for r in challenging) or "—")

    _divider(doc)
    _heading(doc, f"👨 {n1} — Chart Overview", 1)
    _add_overview(doc, chart1, place1)
    _divider(doc)
    _heading(doc, f"👩 {n2} — Chart Overview", 1)
    _add_overview(doc, chart2, place2)

    doc.add_page_break()
    _heading(doc, f"🪐 {n1} — Planetary Positions", 1)
    _add_planets(doc, chart1)
    _heading(doc, f"🪐 {n2} — Planetary Positions", 1)
    _add_planets(doc, chart2)

    doc.add_page_break()
    _heading(doc, f"✨ {n1} — Yogas", 1)
    _add_yogas(doc, chart1)
    _heading(doc, f"✨ {n2} — Yogas", 1)
    _add_yogas(doc, chart2)

    doc.add_page_break()
    _heading(doc, f"🗓️ {n1} — Dasha Sequence", 1)
    _add_dasha(doc, chart1)
    _heading(doc, f"🗓️ {n2} — Dasha Sequence", 1)
    _add_dasha(doc, chart2)

    doc.add_page_break()
    _heading(doc, "📈 Combined Dampati Scores", 1)
    _add_year_scores(doc, combined_scores)
    _divider(doc)
    _heading(doc, f"👨 {n1} Individual Forecast", 2)
    _add_year_scores(doc, s1_data)
    _divider(doc)
    _heading(doc, f"👩 {n2} Individual Forecast", 2)
    _add_year_scores(doc, s2_data)

    doc.add_page_break()
    _add_combined_protection(doc, combined_scores, n1, n2)

    doc.add_page_break()
    _add_predictions(doc, preds1, n1)
    doc.add_page_break()
    _add_predictions(doc, preds2, n2)

    _divider(doc)
    _para(doc, "ॐ तत्सत् · Dampati Jyotisha Report. For spiritual guidance and self-reflection. "
               "Per Yoga Vasishtha: 'The destinies of two souls bound in dharmic union uplift each other.' "
               "Please consult a qualified Jyotishi for major life decisions.",
          italic=True, color=MUTED, size=9)
    return _to_bytes(doc)
